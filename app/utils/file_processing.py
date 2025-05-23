from docx import Document
import fitz  # PyMuPDF
import logging
import re
import os
#import win32com.client as win32
from langchain_community.embeddings import HuggingFaceEmbeddings
import datetime
import subprocess
import openpyxl
from sentence_transformers import SentenceTransformer
import numpy as np
import platform

if platform.system() == "Windows":
    import win32com.client as win32


logger = logging.getLogger(__name__)


# Fonction pour diviser le texte en morceaux de taille minimale avec chevauchement
def divide_text_by_min_words(text, min_words=250, overlap=50):
    """
    Divise le texte en morceaux contenant au moins un nombre minimum de mots,
    avec un chevauchement pour conserver le contexte.
    
    :param text: Le texte à diviser.
    :param min_words: Le nombre minimum de mots par chunk.
    :param overlap: Le nombre de mots à chevaucher entre les chunks.
    :return: Une liste de morceaux de texte.
    """
    words = text.split()
    chunks = []
    i = 0

    while i < len(words):
        end = i + min_words
        chunk = words[i:end]
        chunks.append(' '.join(chunk))
        i += min_words - overlap

    return chunks

# Fonction pour vectoriser les chunks de texte
def vectorize_chunks(chunks):
    """
    Vectorise les chunks de texte en utilisant langchain.
    
    :parametre chunks: Une liste de morceaux de texte.
    :return: Une liste de vecteurs.
    """
    model_name = "sentence-transformers/all-mpnet-base-v2"
    embeddings = HuggingFaceEmbeddings(model_name=model_name)
    vectors = embeddings.embed_documents(chunks)
    return vectors

# Fonction pour nettoyer le texte et supprimer les séquences inutiles
def filter_unnecessary_sequences(text):
    """
    Filtre les séquences inutiles du texte.
    
    :param text: Le texte à filtrer.
    :return: Le texte filtré.
    """
    # Exemple de filtre pour les séquences de points
    text = re.sub(r'\.{3,}', '', text)
    return text

# Fonction pour traiter les fichiers Word
def process_word_file(file_path):
    """
    Traite un fichier Word et le divise en morceaux de texte basés sur les paragraphes.
    Détecte également les images et les inclut dans les chunks.

    """
    try:
        doc = Document(file_path)
        full_text = []

        for para in doc.paragraphs:
            if para.text.strip() == "":
                continue
            full_text.append(para.text)

        # Récupérer le texte des tableaux
        for table_index, table in enumerate(doc.tables):
            table_content = []
            for row in table.rows:
                row_content = []
                for cell in row.cells:
                    cell_text = cell.text.strip()
                    if cell_text:
                        row_content.append(cell_text)
                if row_content:
                    table_content.append(' | '.join(row_content))
            if table_content:
                full_text.append(f"Table {table_index + 1}: " + ' || '.join(table_content))
                print(f"Table {table_index + 1}:")
                for row in table_content:
                    print(row)
        
        
        # Détecter les images
        for rel in doc.part.rels.values():
            if "image" in rel.target_ref:
                full_text.append(f"[Image: {rel.target_ref}]")
        
        text = ' '.join(full_text)
        text = filter_unnecessary_sequences(text)
        words = text.split()
        chunks = []
        current_chunk = []
        overlap = 50
        timestamp = datetime.datetime.now().isoformat()
        file_name = os.path.basename(file_path)
        chunks_id_counter = 1

        for word in words:
            current_chunk.append(word)
            if len(current_chunk) >= 100:
                chunks.append({
                    "file_name": file_name,
                    "pages": "Non trouvé",  # Par défaut pour les fichiers Word
                    "timestamp": timestamp,
                    "chunk_id": chunks_id_counter,
                    "chunk": ' '.join(current_chunk)
                })
                chunks_id_counter += 1
                current_chunk = current_chunk[-overlap:]  # Conserver les 50 derniers mots

        # Ajouter le dernier chunk s'il contient des mots
        if current_chunk:
            chunks.append({
                "file_name": file_name,
                "pages": "Non trouvé",  # Par défaut pour les fichiers Word
                "timestamp": timestamp,
                "chunk_id": chunks_id_counter,
                "chunk": ' '.join(current_chunk)
            })

        # Vectoriser les chunks et ajouter les vecteurs aux chunks
        chunk_texts = [chunk["chunk"] for chunk in chunks]
        vectors = vectorize_chunks(chunk_texts)
        
        for i, chunk in enumerate(chunks):
            chunk["vector"] = vectors[i]

        return chunks
    except Exception as e:
        logger.error(f"Error processing Word file: {e}")
        raise

# Fonction pour traiter les fichiers .doc
def process_doc_file(file_path):
    """
    Traite un fichier Word .doc en le convertissant en .docx,
    puis le divise en morceaux de texte.
    """
    try:
        logger.info(f"Starting to process DOC file: {file_path}")
        
        # Vérifier que le fichier existe
        if not os.path.exists(file_path):
            logger.error(f"File not found: {file_path}")
            raise FileNotFoundError(f"File not found: {file_path}")
        
        # Utiliser un chemin absolu
        abs_file_path = os.path.abspath(file_path)
        
        # Convertir le fichier .doc en .docx en utilisant pywin32
        word = win32.Dispatch("Word.Application")
        doc = word.Documents.Open(abs_file_path)
        docx_file_path = os.path.splitext(abs_file_path)[0] + '.docx'
        doc.SaveAs(docx_file_path, FileFormat=16)  # wdFormatXMLDocument = 16
        doc.Close()
        word.Quit()
        
        # Vérifier que le fichier .docx a été créé
        if not os.path.exists(docx_file_path):
            logger.error(f"Conversion failed, DOCX file not found: {docx_file_path}")
            raise FileNotFoundError(f"Conversion failed, DOCX file not found: {docx_file_path}")
        
        # Traiter le fichier .docx comme un fichier Word
        return process_word_file(docx_file_path)
    except Exception as e:
        logger.error(f"Error processing DOC file: {e}")
        raise

# Fonction pour traiter les fichiers PDF
def process_pdf_file(file_path):
    """
    Traite un fichier PDF et le divise en morceaux de texte basés sur les paragraphes.
    """
    try:
        doc = fitz.open(file_path)
        full_text = []
        current_page = 1
        current_chunk = []
        chunks = []
        overlap = 50
        timestamp = datetime.datetime.now().isoformat()
        file_name = os.path.basename(file_path)
        chunks_id_counter = 1

        for page_num, page in enumerate(doc, start=1):
            text = page.get_text("text")
            text = filter_unnecessary_sequences(text)
            words = text.split()
            
            for word in words:
                current_chunk.append(word)
                if len(current_chunk) >= 100:
                    chunks.append({
                        "file_name": file_name,
                        "pages": f"{current_page}-{page_num}" if current_page != page_num else str(page_num),
                        "timestamp": timestamp,
                        "chunk_id": chunks_id_counter,
                        "chunk": ' '.join(current_chunk)
                    })
                    chunks_id_counter += 1
                    current_chunk = current_chunk[-overlap:]  # Conserver les 50 derniers mots
                    current_page = page_num

        # Ajouter le dernier chunk s'il contient des mots
        if current_chunk:
            chunks.append({
                "file_name": file_name,
                "pages": f"{current_page}-{page_num}" if current_page != page_num else str(page_num),
                "timestamp": timestamp,
                "chunk_id": chunks_id_counter,
                "chunk": ' '.join(current_chunk)
            })

        # Vectoriser les chunks et ajouter les vecteurs aux chunks
        chunk_texts = [chunk["chunk"] for chunk in chunks]
        vectors = vectorize_chunks(chunk_texts)
        
        for i, chunk in enumerate(chunks):
            chunk["vector"] = vectors[i]

        return chunks
    except Exception as e:
        logger.error(f"Error processing PDF file: {e}")
        raise

# Fonction pour charger un fichier Excel avec openpyxl
def load_excel_with_hidden_content(file_name):
    """
    Charge un fichier Excel en incluant les contenus masqués.
    
    Args:
        file_name (str): Nom du fichier Excel (ex. : "example.xlsm").

    Returns:
        dict: Dictionnaire contenant les feuilles et leurs données (visibles et masquées).
    """
    try:
        # Construire le chemin complet
        file_path = file_name
        # Charger le fichier Excel avec openpyxl
        workbook = openpyxl.load_workbook(file_path, data_only=True, keep_links=True)

        data = {}
        for sheet_name in workbook.sheetnames:
            sheet = workbook[sheet_name]
            sheet_data = []
            for row in sheet.iter_rows(values_only=True):
                sheet_data.append(row)
            data[sheet_name] = sheet_data

        return data
    except Exception as e:
        raise ValueError(f"Erreur lors du chargement du fichier Excel : {e}")
    
# Fonction pour nettoyer les données extraites sur excel
def clean_data(data):
    """
    Nettoie les données extraites, supprime les lignes vides et normalise.
    """
    cleaned_data = []
    for row in data:
        if any(row):  # Conserver les lignes qui ne sont pas entièrement vides
            cleaned_data.append([str(cell).strip() if cell else "" for cell in row])
    return cleaned_data

# Fonction pour vectoriser les données sur excel (chunk)
def vectorize_page(page_data):
    """
    Vectorise les données d'une page (chunk).
    """
    text = " ".join([" ".join(map(str, row)) for row in page_data])
    model_name = "sentence-transformers/all-mpnet-base-v2"
    embeddings = HuggingFaceEmbeddings(model_name=model_name)
    embedding = embeddings.embed_documents([text])
    return embedding.tolist()  # Convertir en liste Python

# Fonction pour traiter les fichiers Excel
def process_excel_file(file_name):
    """
    Traite un fichier Excel et génère un chunk par page (feuille).
    """
    try:
        # Charger tout le contenu, y compris les données masquées
        data = load_excel_with_hidden_content(file_name)

        chunks = []
        timestamp = datetime.datetime.now().isoformat()
        file_name = os.path.basename(file_name)
        chunks_id_counter = 1

        for sheet_name, sheet_data in data.items():
            # Nettoyer les données
            cleaned_data = clean_data(sheet_data)
            
            # Vectoriser les données de la feuille entière
            #embedding = vectorize_page(cleaned_data)

            chunk_text = ' '.join([' '.join(map(str, row)) for row in cleaned_data])
            

            # Ajouter les résultats
            chunks.append({
                "file_name": file_name,
                "pages": sheet_name,
                "timestamp": timestamp,
                "chunk_id": chunks_id_counter,
                "chunk": chunk_text,
                
            })
            chunks_id_counter += 1

        chunk_texts = [chunk["chunk"] for chunk in chunks]
        vectors = vectorize_chunks(chunk_texts)
        
        for i, chunk in enumerate(chunks):
            chunk["vector"] = vectors[i]

        return chunks
    except Exception as e:
        return {"status": "error", "message": str(e)}

# Fonction principale pour traiter les fichiers
def process_file(file_path):
    """
    Traite un fichier en fonction de son type et le divise en morceaux de texte.
    
    :param file_path: Le chemin du fichier.
    :return: Une liste de morceaux de texte avec des informations sur le fichier.
    """
    if file_path.endswith('.pdf'):
        return process_pdf_file(file_path)
    elif file_path.endswith('.docx'):
        return process_word_file(file_path)
    elif file_path.endswith('.doc'):
        return process_doc_file(file_path)
    elif file_path.endswith('.xlsx') or file_path.endswith('.xlsm'):
        return process_excel_file(file_path)
    
    else:
        raise ValueError("Unsupported file type")
